"""Redact PII in a Word document while preserving most Word formatting.

Usage:
    python redact.py "Red Herring Prospectus.docx" redacted_prospectus.docx

The detector is deliberately conservative for generic dates: a date is redacted
only when its surrounding paragraph says it is a date of birth.  This avoids
destroying dates that are material to a prospectus.
"""

from __future__ import annotations

import argparse
import json
import os
import re
import zipfile
from collections import Counter
from dataclasses import dataclass, field
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.document import Document as _Document
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph


# Patterns with a low false-positive rate.  Phone detection has an additional
# validation step so identifiers, financial values, and years are not masked.
EMAIL = re.compile(r"\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}\b")
IP_ADDRESS = re.compile(r"(?<![\d.])(?:25[0-5]|2[0-4]\d|1?\d?\d)(?:\.(?:25[0-5]|2[0-4]\d|1?\d?\d)){3}(?![\d.])")
SSN = re.compile(r"\b(?!000|666|9\d\d)\d{3}[- ]?(?!00)\d{2}[- ]?(?!0000)\d{4}\b")
CARD = re.compile(r"(?<!\d)(?:\d[ -]?){13,19}(?!\d)")
PHONE = re.compile(
    r"(?<![\dA-Za-z])(?:\+\s?91[\s.-]*)?(?:\(\d{2,5}\)|\d{2,5})(?:[\s.-]+\d{2,10}){1,3}(?![\dA-Za-z])"
)
DATE = re.compile(r"\b(?:January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{1,2},\s+\d{4}\b", re.I)

# Addresses and organisation names are necessarily heuristic.  These patterns
# are intentionally restricted to common business-document forms.
ADDRESS = re.compile(
    r"\b(?:\d{1,5}[/-]\d{0,5}|\d{1,5})\s*,?\s*(?:[A-Za-z][\w.'-]*\s*){0,12}"
    r"(?:Road|Rd\.?|Street|St\.?|Avenue|Ave\.?|Lane|Ln\.?|Marg|Nagar|"
    r"Village|Taluka|Tower|Floor|Centre|Center|Building|Complex|Park)\b"
    r"[^;\n]{0,160}?\b\d{6}\b(?:\s*,?\s*India)?",
    re.I,
)
COMPANY = re.compile(
    r"\b(?:[A-Z][A-Za-z&.'-]*\s+){0,8}"
    r"(?:Private Limited|Public Limited|Limited|Ltd\.?|LLP|L\.L\.P\.|"
    r"Inc\.?|Corporation|Corp\.?|Bank|Trust)\b"
)
UPPERCASE_COMPANY = re.compile(
    r"\b(?:[A-Z]{2,}\s+){1,8}(?:LIMITED|LTD\.?|LLP|INC\.?|CORPORATION|CORP\.?|BANK|TRUST)\b"
)
TITLE_NAME = re.compile(r"\b(?:Mr\.?|Mrs\.?|Ms\.?|Dr\.?)\s+([A-Z][a-z]+(?:\s+[A-Z][a-z]+){1,4})\b")
CONTACT_NAME = re.compile(
    r"(?i)(?:contact\s*person|authori[sz]ed\s*signatory)\s*:\s*"
    r"([A-Z][A-Za-z.'-]+(?:\s*/\s*[A-Z][A-Za-z.'-]+|\s+[A-Z][A-Za-z.'-]+){1,5})"
)


def iter_block_items(parent: _Document | _Cell) -> Iterable[Paragraph | Table]:
    """Yield paragraphs/tables in document order, including nested tables."""
    if isinstance(parent, _Document):
        element = parent.element.body
    elif isinstance(parent, _Cell):
        element = parent._tc
    else:  # Header/footer containers expose the same XML child structure.
        element = parent._element
    for child in element.iterchildren():
        if child.tag.endswith("}p"):
            yield Paragraph(child, parent)
        elif child.tag.endswith("}tbl"):
            yield Table(child, parent)


def iter_paragraphs(parent: _Document | _Cell) -> Iterable[Paragraph]:
    for block in iter_block_items(parent):
        if isinstance(block, Paragraph):
            yield block
        else:
            for row in block.rows:
                for cell in row.cells:
                    yield from iter_paragraphs(cell)


def document_paragraphs(doc: _Document) -> Iterable[Paragraph]:
    """Include body, table cells, headers, and footers exactly once each."""
    yield from iter_paragraphs(doc)
    seen_parts: set[str] = set()
    for section in doc.sections:
        for part in (section.header, section.footer, section.first_page_header,
                     section.first_page_footer, section.even_page_header,
                     section.even_page_footer):
            part_id = str(part.part.partname)
            if part_id not in seen_parts:
                seen_parts.add(part_id)
                yield from iter_paragraphs(part)


def luhn_valid(value: str) -> bool:
    digits = [int(x) for x in re.sub(r"\D", "", value)]
    if not 13 <= len(digits) <= 19 or len(set(digits)) == 1:
        return False
    total = 0
    for i, digit in enumerate(reversed(digits)):
        total += digit if i % 2 == 0 else (digit * 2 - 9 if digit > 4 else digit * 2)
    return total % 10 == 0


def plausible_phone(value: str) -> bool:
    digits = re.sub(r"\D", "", value)
    # Indian contact numbers in this prospectus include +91 / 91 prefixes and
    # landlines. This excludes unseparated dates, IDs, and currency amounts.
    return (len(digits) in (10, 11, 12) and (" " in value or "-" in value or "+" in value or "(" in value)
            and not re.fullmatch(r"\d{4}-\d{4}", value.strip()))


@dataclass
class ReplacementStore:
    maps: dict[str, dict[str, str]] = field(default_factory=lambda: {})
    counts: Counter = field(default_factory=Counter)

    def replacement(self, kind: str, original: str) -> str:
        lookup = self.maps.setdefault(kind, {})
        key = original.casefold() if kind in {"email", "company", "person", "address"} else original
        if key not in lookup:
            index = len(lookup) + 1
            lookup[key] = self._fake(kind, index, original)
        self.counts[kind] += 1
        return lookup[key]

    @staticmethod
    def _fake(kind: str, index: int, original: str) -> str:
        if kind == "email":
            return f"contact{index:03d}@example.test"
        if kind == "phone":
            return f"+91 70000 {index:05d}"
        if kind == "ip":
            return f"203.0.113.{index % 254 + 1}"
        if kind == "ssn":
            return f"900-10-{index:04d}"
        if kind == "card":
            return f"4111 1111 1111 {index:04d}"
        if kind == "dob":
            return f"January {index:02d}, 1970"
        if kind == "address":
            return f"{index} Example Avenue, Sample City - 000000, India"
        if kind == "company":
            return f"Example Organisation {index} Limited"
        if kind == "person":
            first = ("Aarav", "Diya", "Kabir", "Meera", "Ishaan", "Anaya")[index % 6]
            last = ("Shah", "Rao", "Kapoor", "Iyer", "Malik", "Sen")[index % 6]
            return f"{first} {last}"
        raise ValueError(kind)


def find_matches(
    text: str, known_names: tuple[str, ...] = (), known_addresses: tuple[str, ...] = ()
) -> list[tuple[int, int, str, str]]:
    """Return non-overlapping spans, prioritising structured PII over names."""
    candidates: list[tuple[int, int, str, str]] = []

    def add(pattern: re.Pattern, kind: str, predicate=lambda _: True, group: int = 0):
        for match in pattern.finditer(text):
            value = match.group(group)
            if predicate(value):
                start, end = match.span(group)
                candidates.append((start, end, kind, value))

    add(EMAIL, "email")
    add(IP_ADDRESS, "ip")
    add(SSN, "ssn")
    # Gate the broad expressions. This makes large financial tables fast while
    # retaining coverage for text that can actually contain the category.
    if re.search(r"\d(?:[ -]?\d){12,}", text):
        add(CARD, "card", luhn_valid)
    if any(mark in text for mark in ("+", "(", "-", " ")):
        add(PHONE, "phone", plausible_phone)
    lower = text.lower()
    address_words = ("road", "street", "avenue", "marg", "nagar", "village", "taluka", "tower", "floor", "centre", "center", "building", "complex", "park")
    if re.search(r"\b\d{3}\s?\d{3}\b", text) and any(word in lower for word in address_words):
        add(ADDRESS, "address")
    company_words = ("limited", "ltd", "llp", "inc", "corporation", "corp", "bank", "trust")
    if any(word in lower for word in company_words):
        add(COMPANY, "company", lambda x: len(x.split()) >= 2)
    add(UPPERCASE_COMPANY, "company")
    add(TITLE_NAME, "person", group=1)
    add(CONTACT_NAME, "person", group=1)
    for name in known_names:
        for match in re.finditer(r"(?<!\w)" + re.escape(name) + r"(?!\w)", text, re.I):
            candidates.append((match.start(), match.end(), "person", match.group()))
    for address in known_addresses:
        for match in re.finditer(re.escape(address), text, re.I):
            candidates.append((match.start(), match.end(), "address", match.group()))
    if re.search(r"(?i)\b(?:date\s+of\s+birth|dob)\b", text):
        add(DATE, "dob")

    # Longest/highest-priority match wins where heuristic categories overlap.
    priority = {"email": 0, "ip": 1, "ssn": 2, "card": 3, "phone": 4,
                "address": 5, "person": 6, "company": 7, "dob": 8}
    selected: list[tuple[int, int, str, str]] = []
    for item in sorted(candidates, key=lambda x: (x[0], priority[x[2]], -(x[1] - x[0]))):
        if not any(item[0] < chosen[1] and chosen[0] < item[1] for chosen in selected):
            selected.append(item)
    return selected


def replace_span(paragraph: Paragraph, start: int, end: int, replacement: str) -> None:
    """Replace a character span even if Word split it across formatted runs."""
    runs = paragraph.runs
    positions: list[tuple[int, int]] = []
    cursor = 0
    for run in runs:
        positions.append((cursor, cursor + len(run.text)))
        cursor += len(run.text)
    if start == end or not runs:
        return
    start_run = next(i for i, (a, b) in enumerate(positions) if a <= start < b)
    end_run = next(i for i, (a, b) in enumerate(positions) if a < end <= b)
    a0, _ = positions[start_run]
    a1, _ = positions[end_run]
    if start_run == end_run:
        value = runs[start_run].text
        runs[start_run].text = value[: start - a0] + replacement + value[end - a0 :]
        return
    runs[start_run].text = runs[start_run].text[: start - a0] + replacement
    for index in range(start_run + 1, end_run):
        runs[index].text = ""
    runs[end_run].text = runs[end_run].text[end - a1 :]


def redact_paragraph(
    paragraph: Paragraph, store: ReplacementStore, known_names: tuple[str, ...], known_addresses: tuple[str, ...]
) -> None:
    text = "".join(run.text for run in paragraph.runs)
    # Work right-to-left so offsets remain valid after each replacement.
    for start, end, kind, value in reversed(find_matches(text, known_names, known_addresses)):
        replace_span(paragraph, start, end, store.replacement(kind, value))


def scrub_hyperlink_targets(path: Path, store: ReplacementStore) -> None:
    """Remove PII left in Word field instructions such as ``mailto:...``.

    python-docx exposes the displayed field result as a run but not the field
    instruction.  A post-save pass is therefore needed to prevent the original
    e-mail address or telephone number being recoverable from DOCX XML.
    """
    target = path.with_name(path.stem + ".scrubbing" + path.suffix)
    with zipfile.ZipFile(path, "r") as source, zipfile.ZipFile(target, "w", zipfile.ZIP_DEFLATED) as clean:
        for info in source.infolist():
            payload = source.read(info.filename)
            if info.filename.startswith("word/") and info.filename.endswith(".xml"):
                xml = payload.decode("utf-8")
                xml = re.sub(
                    r"(?i)(mailto:)([A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,})",
                    lambda match: match.group(1) + store.replacement("email", match.group(2)),
                    xml,
                )

                def replace_tel(match: re.Match) -> str:
                    number = match.group(2)
                    digits = re.sub(r"\D", "", number)
                    # Field targets do not always use the displayed grouping.
                    if 10 <= len(digits) <= 12:
                        return match.group(1) + store.replacement("phone", number)
                    return match.group(0)

                xml = re.sub(r"(?i)(tel:)([+0-9(). -]{8,22})", replace_tel, xml)
                payload = xml.encode("utf-8")
            clean.writestr(info, payload)
    os.replace(target, path)


def extract_table_names(doc: _Document) -> tuple[str, ...]:
    """Learn high-confidence personal names from clearly labelled people tables.

    This avoids treating every title-cased phrase in narrative prose as a name.
    It is useful for prospectuses, CVs, and registers that have a Name column.
    """
    names: set[str] = set()
    for table in doc.tables:
        if len(table.rows) < 2:
            continue
        headings = " ".join(cell.text.lower() for cell in table.rows[0].cells)
        people_table = "name" in headings and any(word in headings for word in (
            "director", "promoter", "designation", "address", "age", "qualification"))
        if not people_table:
            continue
        for row in table.rows[1:]:
            value = re.sub(r"[*\u2020\u2021]+$", "", row.cells[0].text.strip())
            if re.fullmatch(r"[A-Z][A-Za-z.'-]+(?:\s+[A-Z][A-Za-z.'-]+){1,4}", value):
                names.add(value)
    return tuple(sorted(names, key=len, reverse=True))


def extract_table_addresses(doc: _Document) -> tuple[str, ...]:
    """Learn complete addresses from tables with an explicit Address column."""
    addresses: set[str] = set()
    for table in doc.tables:
        if len(table.rows) < 2:
            continue
        columns = [cell.text.strip().casefold() for cell in table.rows[0].cells]
        address_columns = [i for i, heading in enumerate(columns) if "address" in heading]
        for row in table.rows[1:]:
            for index in address_columns:
                if index < len(row.cells):
                    value = " ".join(row.cells[index].text.split())
                    if len(value) >= 15:
                        addresses.add(value)
    return tuple(sorted(addresses, key=len, reverse=True))


def extract_labelled_addresses(doc: _Document) -> tuple[str, ...]:
    """Learn complete addresses following explicit Office/Address labels."""
    addresses: set[str] = set()
    label = re.compile(r"(?i)\b(?:registered\s+office|corporate\s+office|mailing\s+address|address)\s*:\s*(.+)")
    for paragraph in document_paragraphs(doc):
        text = "".join(run.text for run in paragraph.runs).strip()
        match = label.search(text)
        if not match:
            continue
        value = match.group(1).split(";")[0].strip()
        has_postcode = bool(re.search(r"\b\d{3}\s?\d{3}\b", value))
        if has_postcode and len(value) >= 15:
            addresses.add(value)
    return tuple(sorted(addresses, key=len, reverse=True))


def redact(input_path: Path, output_path: Path, report_path: Path | None = None) -> dict:
    doc = Document(input_path)
    store = ReplacementStore()
    known_names = extract_table_names(doc)
    known_addresses = tuple(sorted(
        set(extract_table_addresses(doc)) | set(extract_labelled_addresses(doc)), key=len, reverse=True
    ))
    for paragraph in document_paragraphs(doc):
        redact_paragraph(paragraph, store, known_names, known_addresses)
    doc.save(output_path)
    scrub_hyperlink_targets(output_path, store)
    report = {
        "input": str(input_path), "output": str(output_path),
        "occurrences_redacted": dict(store.counts),
        "unique_values_redacted": {kind: len(values) for kind, values in store.maps.items()},
    }
    if report_path:
        report_path.write_text(json.dumps(report, indent=2), encoding="utf-8")
    return report


def main() -> None:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("input", type=Path)
    parser.add_argument("output", type=Path)
    parser.add_argument("--report", type=Path, help="Optional JSON redaction-count report")
    args = parser.parse_args()
    print(json.dumps(redact(args.input, args.output, args.report), indent=2))


if __name__ == "__main__":
    main()
